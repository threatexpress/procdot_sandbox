import csv
import sys
import re
from docx import Document

############################################
# Reference
############################################
# CSV Parsing - http://pymotw.com/2/csv/

############################################
# Initialize Variables
############################################
# Open CSV and files
f = open(sys.argv[1], 'rt')
logText = open('log.txt','w')
logWord = 'log.docx'
document = Document()

# Row Count
totalrows = 0
# priorRow Holder
priorRow = {}
priorRow['\xef\xbb\xbf"Time of Day"'] = ''
priorRow['Process Name']  = ''
priorRow['PID']  = ''
priorRow['Operation']  = ''
priorRow['Path']  = ''
priorRow['Result']  = ''
priorRow['Detail']  = ''
priorRow['TID']  = ''

# Open Config files
processIgnore             = [line.strip() for line in open('./config/processIgnore.txt')]
fullFilePathIgnore        = [line.strip() for line in open('./config/fullFilePathIgnore.txt')]
fullPathRegistryKeyIgnore = [line.strip() for line in open('./config/fullPathRegistryKeyIgnore.txt')]
regexFilePathIgnore       = [line.strip() for line in open('./config/regexFilePathIgnore.txt')]
regexRegistryKeyIgnore    = [line.strip() for line in open('./config/regexRegistryKeyIgnore.txt')]

# Parser Counters
fullFilePathIgnoreCOUNT = 0
fullPathRegistryKeyIgnoreCOUNT = 0
regexFilePathIgnoreCOUNT = 0
regexRegistryKeyIgnoreCOUNT = 0

############################################
# Functions
############################################
def printStatus(i):
    ''' Print current progress '''
    print "Progress: {0}% ({1} of {2})".format(int(float(i) / float(totalrows) * 100),i,totalrows)

def updatepriorRow(row):
    ''' Store the current row for repetition comparison '''
    
    #priorRow['\xef\xbb\xbf"Time of Day"'] = row['\xef\xbb\xbf"Time of Day"']
    priorRow['Process Name'] = row['Process Name'] 
    priorRow['PID'] = row['PID']
    priorRow['Operation'] = row['Operation']
    priorRow['Path'] = row['Path']
    priorRow['Result'] = row['Result']
    priorRow['Detail'] = ['Detail']
    priorRow['TID'] = ['TID']

def ignoreRepeat(row):
    ''' Check to see if the prior Operation is the same as the current.
        Do not logText if it was the same.'''
    #logText.write("{} {} {} {}".format(priorRow['Operation'], row['Operation'], priorRow['Path'], row['Path'])+ "\n")
    if (priorRow['Operation'] == row['Operation']) and (priorRow['Path'] == row['Path']):
        return True
    else:
        return False

def ignoreProcess(processName):
    ''' Compares process with a list of process to ignore and not logText. '''

    if processName in processIgnore:
        return True
    else:
        return False

def ignoreFullFilePath(path):
    ''' Compares file path with a list of file paths to ignore and not logText. '''
    global fullFilePathIgnoreCOUNT

    if path.lower() in fullFilePathIgnore:
        fullFilePathIgnoreCOUNT += 1
        return True
    else:
        return False

def ignoreRegexFilePath(path):
    ''' Compares file path using regex with a list of file paths to ignore and not logText. '''
    global regexFilePathIgnoreCOUNT

    pathFound = False
    for search in regexFilePathIgnore:
        p = re.compile(search, re.IGNORECASE)
        found = p.match(path.lower())
        if found:
            pathFound = True
            break

    if pathFound:  
        regexFilePathIgnoreCOUNT += 1
        return True
    else:
        return False    

def ignoreFullPathRegistryKey(path): 
    ''' Compares registy path with a list of paths to ignore and not logText. '''
    global fullPathRegistryKeyIgnoreCOUNT

    if path.lower() in fullPathRegistryKeyIgnore:
        fullPathRegistryKeyIgnoreCOUNT += 1
        return True
    else:
        return False

def ignoreRegexRegistryKey(path):
    ''' Compares registry key using regex with a list of key values to ignore and not logText. '''
    global regexRegistryKeyIgnoreCOUNT

    pathFound = False
    for search in regexRegistryKeyIgnore:
        p = re.compile(search, re.IGNORECASE)       

        found = p.match(path.lower())
        if found:
            pathFound = True
            break

    if pathFound:
        regexRegistryKeyIgnoreCOUNT += 1
        return True
    else:
        return False  

def processRow(row):
    ''' Process row for data 
        Row Header Reference
        "Time of Day","Process Name","PID","Operation","Path","Result","Detail","TID"
    '''
    # Capture row values
    #TOD = row['\xef\xbb\xbf"Time of Day"']
    PRN = row['Process Name']
    PID = row['PID']
    OPN = row['Operation']
    PTH = row['Path']
    RST = row['Result']
    DTL = row['Detail']
    TID = row['TID']

    # Check for rows to Ignore - START
    if ignoreProcess(PRN):
        return

    if ignoreFullFilePath(PTH):
        return

    if ignoreFullPathRegistryKey(PTH):
        return

    if ignoreRegexFilePath(PTH):
        return

    if ignoreRegexRegistryKey(PTH):
        return
    # Check for rows to Ignore - END

    elif OPN == "WriteFile":
        if ignoreRepeat(row):
            return
        else:            
            line = "Thread {0} of process {1} (PID: {2}) wrote to file {3}.".format(TID, PRN, PID, PTH) 
            logText.write(line + "\n")
            line = "Thread {0} of process {1} (PID: {2}) wrote to file.".format(TID, PRN, PID) 
            document.add_paragraph(line)
            #document.add_paragraph("{0}ACTION: {1}".format("\t",OPN))
            #document.add_paragraph("{0}THREAD: {1}".format("\t\t",TID))
            #document.add_paragraph("{0}PATH: {1}".format("\t\t",PTH))
            #document.add_paragraph("{0}VALUE: {1}".format("\t\t",DTL))
            updatepriorRow(row)
        return

    elif OPN == "RegQueryValue":
        if ignoreRepeat(row):            
            return
        else:
            line = "Thread {0} of process {1} (PID: {2}) queried registy key {3}.  Result: {4}".format(TID, PRN, PID, PTH, DTL)
            logText.write(line + "\n")
            line = "Thread {0} of process {1} (PID: {2}) queried registy key.".format(TID, PRN, PID)
            #document.add_paragraph(line)
            #document.add_paragraph("{0}ACTION: {1}".format("\t",OPN))
            #document.add_paragraph("{0}THREAD: {1}".format("\t\t",TID))
            #ocument.add_paragraph("{0}PATH: {1}".format("\t\t",PTH))
            #ocument.add_paragraph("{0}VALUE: {1}".format("\t\t",DTL))            
            updatepriorRow(row)
        return

    elif OPN == "RegSetValue":
        if ignoreRepeat(row):            
            return
        else:
            line = "Thread {0} of process {1} (PID: {2}) set registy key {3}.  Result: {4}".format(TID, PRN, PID, PTH, DTL)
            logText.write(line + "\n")
            line = "Thread {0} of process {1} (PID: {2}) set registy key.".format(TID, PRN, PID)
            #document.add_paragraph(line)
            #document.add_paragraph("{0}ACTION: {1}".format("\t",OPN))
            #ocument.add_paragraph("{0}THREAD: {1}".format("\t\t",TID))
            #document.add_paragraph("{0}PATH: {1}".format("\t\t",PTH))
            #document.add_paragraph("{0}VALUE: {1}".format("\t\t",DTL))            
            updatepriorRow(row)
        return
  
    elif OPN == "SetDispositionInformationFile":
        if ignoreRepeat(row):            
            return
        else:
            line = "Thread {0} of process {1} (PID: {2}) deleted file {3}.  Result: {4}".format(TID, PRN, PID, PTH, DTL)
            logText.write(line + "\n")
            line = "Thread {0} of process {1} (PID: {2}) deleted file.".format(TID, PRN, PID)
            #document.add_paragraph(line)
            #document.add_paragraph("{0}ACTION: {1}".format("\t",OPN))
            #document.add_paragraph("{0}THREAD: {1}".format("\t\t",TID))
            #document.add_paragraph("{0}PATH: {1}".format("\t\t",PTH))
            #document.add_paragraph("{0}VALUE: {1}".format("\t\t",DTL))            
            #document.add_paragraph(line)
            updatepriorRow(row)
        return
    else:
        # Ignore Row
        return
############################################
# Main
############################################
try:

    # Count Rows        
    readerCntr = csv.DictReader(f)
    print "Reading CSV..."    
    for row in readerCntr:
        totalrows += 1
    f.seek(0)  # Reset rile position

    # Process Rows
    print "Processing CSV..."
    reader = csv.DictReader(f)
    i = 1
    for row in reader:       
        if i%10000 == 0:
            printStatus(i)
        processRow(row)
        
        i += 1  
    print ""
    print "--------------------------------------------------------" 
    print "                         Summary"
    print "--------------------------------------------------------"    
    print "Total Items Processed: {0}".format(i,totalrows)
    print "--------------------------------------------------------"
    print "Items Filtered"
    print "---------------"
    print " fullFilePathIgnore        {0}".format(fullFilePathIgnoreCOUNT)
    print " fullPathRegistryKeyIgnore {0}".format(fullPathRegistryKeyIgnoreCOUNT)
    print " regexFilePathIgnore       {0}".format(regexFilePathIgnoreCOUNT)
    print " regexRegistryKeyIgnore    {0}".format(regexRegistryKeyIgnoreCOUNT)
    print "--------------------------------------------------------"

finally:
    f.close()
    logText.close()
    document.save(logWord)







